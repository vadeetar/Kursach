#!/usr/bin/env python3
"""Генерация пояснительной записки в формате Word (.docx)."""

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    raise SystemExit("Установите: pip install python-docx")

OUTPUT = Path(__file__).resolve().parent.parent / "POYASNITELNA_ZAPISKA.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Cm(1.25)


def add_centered(doc: Document, text: str, size: int = 14, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    set_normal_style(doc)

    # Титул
    add_centered(doc, "Министерство науки и высшего образования Российской Федерации", 12)
    add_centered(doc, "(наименование образовательной организации — уточнить по шаблону вуза)", 12)
    doc.add_paragraph()
    add_centered(doc, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", 16, bold=True)
    add_centered(doc, "к курсовому проекту", 14)
    add_centered(doc, "по дисциплине «Методы и технологии программирования»", 14)
    doc.add_paragraph()
    add_centered(doc, "«Платформа управления уязвимостями»", 14, bold=True)
    add_centered(doc, "(вариант 4)", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Выполнил: студент гр. 221131", indent=False)
    add_para(doc, "Тарасов Вадим Романович", indent=False)
    add_para(doc, "Руководитель: _______________________________", indent=False)
    add_para(doc, "Дата: «___» __________ 2026 г.", indent=False)
    doc.add_page_break()

    # Содержание
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    for item in [
        "1. Введение",
        "2. Аналитическая часть",
        "3. Проектная часть",
        "4. Технологическая часть",
        "5. Реализация",
        "6. Тестирование",
        "7. Заключение",
        "8. Список использованной литературы",
        "Приложения",
    ]:
        add_para(doc, item, indent=False)
    doc.add_page_break()

    # 1. Введение
    add_heading(doc, "1. ВВЕДЕНИЕ", 1)
    add_heading(doc, "Актуальность", 2)
    add_para(
        doc,
        "Прикладное программное обеспечение организаций строится на цепочках зависимостей "
        "(PyPI, npm, Maven и др.). Каждая библиотека может содержать известные уязвимости (CVE). "
        "По данным NIST, ежегодно публикуется более 20 000 новых CVE; значительная доля инцидентов "
        "связана с неустранёнными уязвимостями. Без автоматизированного сбора, приоритизации и учёта "
        "статуса исправления управление рисками становится неэффективным.",
    )
    add_heading(doc, "Цель работы", 2)
    add_para(
        doc,
        "Разработать программный продукт для автоматического сбора данных об уязвимостях из открытых "
        "источников, сопоставления с инвентаризацией ПО, контекстной приоритизации (CVSS, EPSS, "
        "критичность актива) и ведения статуса исправления через веб-интерфейс и REST API.",
    )
    add_heading(doc, "Задачи", 2)
    tasks = [
        "Проанализировать предметную область и аналоги.",
        "Спроектировать архитектуру и модель данных.",
        "Реализовать интеграцию с NVD API и EPSS API.",
        "Реализовать инвентаризацию активов и сопоставление с CVE.",
        "Разработать алгоритм приоритизации и SLA.",
        "Создать веб-интерфейс и документированный REST API (OpenAPI).",
        "Обеспечить качество: pytest, SAST, контейнеризация, CI/CD.",
        "Подготовить документацию и материалы к защите.",
    ]
    for i, t in enumerate(tasks, 1):
        add_para(doc, f"{i}. {t}", indent=False)

    # 2. Аналитическая
    add_heading(doc, "2. АНАЛИТИЧЕСКАЯ ЧАСТЬ", 1)
    add_heading(doc, "2.1. Предметная область", 2)
    add_para(
        doc,
        "Цикл управления уязвимостями включает этапы: обнаружение, оценка, планирование, "
        "исправление и верификация.",
    )
    add_table(
        doc,
        ["Источник", "Назначение"],
        [
            ["NVD (nvd.nist.gov)", "CVE, CVSS 3.x, CPE, описания"],
            ["EPSS (first.org/epss)", "Вероятность появления эксплойта"],
            ["Инвентарь организации", "Пакеты, версии, критичность"],
        ],
    )
    add_heading(doc, "2.2. Анализ аналогов", 2)
    add_table(
        doc,
        ["Продукт", "Ограничение", "Наше решение"],
        [
            ["Snyk, WhiteSource", "Коммерция, облако", "Локальное развёртывание"],
            ["Dependabot", "Только GitHub", "Независимый REST API"],
            ["Ручной аудит", "Медленно, ошибки", "Автоматизация NVD + matcher"],
        ],
    )
    add_heading(doc, "2.3. Требования к системе", 2)
    add_para(doc, "Функциональные требования (реализованы):", bold=True)
    for req in [
        "Синхронизация CVE из NVD API 2.0;",
        "Обновление EPSS для записей в БД;",
        "CRUD инвентаря ПО;",
        "Автоматическое создание findings;",
        "Фильтрация находок и учёт SLA;",
        "Отчёт executive-summary;",
        "Веб-дашборд.",
    ]:
        add_para(doc, f"— {req}", indent=False)

    # 3. Проектная
    add_heading(doc, "3. ПРОЕКТНАЯ ЧАСТЬ", 1)
    add_heading(doc, "3.1. Архитектура", 2)
    add_para(
        doc,
        "Применена трёхуровневая архитектура: уровень представления (веб-UI и OpenAPI), "
        "бизнес-логика (FastAPI, сервисы NVD, EPSS, Matcher, Prioritizer), уровень данных "
        "(SQLAlchemy, PostgreSQL) и внешние API (NVD, EPSS).",
    )
    add_para(doc, "[Рисунок 1 — Архитектура системы. Вставьте скриншот или диаграмму из docs/ARCHITECTURE.md]", bold=True)
    add_heading(doc, "3.2. Модель данных", 2)
    for m in [
        "Vulnerability — запись CVE, CVSS, EPSS, CWE;",
        "Asset — пакет в инвентаре (ecosystem, name, version);",
        "Finding — связь уязвимости и актива, статус, приоритет, SLA;",
        "Remediation — рекомендации по обновлению;",
        "ScanMetadata — журнал синхронизаций.",
    ]:
        add_para(doc, f"— {m}", indent=False)
    add_heading(doc, "3.3. REST API", 2)
    add_table(
        doc,
        ["Метод", "Endpoint", "Назначение"],
        [
            ["GET", "/api/v1/health", "Проверка сервиса"],
            ["GET/POST", "/api/v1/vulnerabilities", "Работа с CVE"],
            ["CRUD", "/api/v1/assets", "Инвентарь ПО"],
            ["GET/PATCH", "/api/v1/findings", "Находки"],
            ["POST", "/api/v1/tasks/sync-nvd", "Синхронизация NVD"],
            ["GET", "/api/v1/reports/executive-summary", "Отчёт"],
        ],
    )

    # 4. Технологическая
    add_heading(doc, "4. ТЕХНОЛОГИЧЕСКАЯ ЧАСТЬ", 1)
    add_table(
        doc,
        ["Компонент", "Технология"],
        [
            ["Язык", "Python 3.12"],
            ["Framework", "FastAPI"],
            ["Валидация", "Pydantic V2"],
            ["СУБД", "PostgreSQL 15"],
            ["Контейнеры", "Docker, Docker Compose"],
            ["Тесты", "pytest"],
            ["CI/CD", "GitHub Actions"],
            ["Репозиторий", "GitHub"],
        ],
    )
    add_para(
        doc,
        "Выбор стека обусловлен требованиями варианта 4 методических указаний: Python, FastAPI, "
        "PostgreSQL, интеграция с NVD API.",
    )

    # 5. Реализация
    add_heading(doc, "5. РЕАЛИЗАЦИЯ", 1)
    add_para(
        doc,
        "Исходный код размещён в репозитории: https://github.com/vadeetar/Kursach. "
        "Ключевые модули: nvd_collector.py, epss_collector.py, matcher.py, prioritizer.py, routes.py.",
    )
    add_para(doc, "[Рисунок 2 — Веб-интерфейс, дашборд. Скриншот http://127.0.0.1:8000]", bold=True)
    add_para(doc, "[Рисунок 3 — Swagger API. Скриншот /api/v1/docs]", bold=True)
    add_para(
        doc,
        "Алгоритм приоритизации учитывает CVSS, EPSS и критичность актива. "
        "SLA: critical — 7 дней, high — 14, medium — 30, low — 90.",
    )

    # 6. Тестирование
    add_heading(doc, "6. ТЕСТИРОВАНИЕ", 1)
    add_table(
        doc,
        ["Уровень", "Содержание"],
        [
            ["Unit", "Модели, prioritizer, парсер NVD"],
            ["Integration", "REST API, БД"],
            ["CI", "pytest, Bandit, Ruff"],
        ],
    )
    add_para(
        doc,
        "Выполнено 24 автоматических теста (pytest). Покрытие кода — около 71%. "
        "Статический анализ Bandit выполняется в GitHub Actions. "
        "[Рисунок 4 — Успешный CI на GitHub]",
        bold=False,
    )

    # 7. Заключение
    add_heading(doc, "7. ЗАКЛЮЧЕНИЕ", 1)
    add_para(
        doc,
        "В ходе выполнения курсового проекта разработана платформа управления уязвимостями, "
        "соответствующая варианту 4. Реализованы сбор CVE из NVD, обновление EPSS, инвентаризация ПО, "
        "сопоставление с уязвимостями, приоритизация и веб-интерфейс для отслеживания статуса исправления. "
        "Проект развёртывается через Docker Compose или скрипт start.ps1, сопровождается тестами и CI/CD.",
    )
    add_table(
        doc,
        ["Требование методички", "Статус"],
        [
            ["Python + FastAPI + PostgreSQL + NVD", "Выполнено"],
            ["Веб-интерфейс remediation", "Выполнено"],
            ["Git, Docker, тесты, SAST", "Выполнено"],
            ["Документация", "Выполнено"],
        ],
    )

    # 8. Литература
    add_heading(doc, "8. СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", 1)
    refs = [
        "NIST National Vulnerability Database [Электронный ресурс]. — URL: https://nvd.nist.gov",
        "FIRST EPSS [Электронный ресурс]. — URL: https://www.first.org/epss",
        "MITRE CVE Program [Электронный ресурс]. — URL: https://www.cve.org",
        "FastAPI Documentation [Электронный ресурс]. — URL: https://fastapi.tiangolo.com",
        "SQLAlchemy 2.0 Documentation [Электронный ресурс]. — URL: https://docs.sqlalchemy.org",
        "Pydantic V2 Documentation [Электронный ресурс]. — URL: https://docs.pydantic.dev",
        "Docker Documentation [Электронный ресурс]. — URL: https://docs.docker.com",
        "GitHub Actions Documentation [Электронный ресурс]. — URL: https://docs.github.com/actions",
        "Макконнелл С. Совершенный код. — Microsoft Press, 2023.",
        "Фаулер М. Архитектура корпоративных приложений. — Addison-Wesley, 2022.",
        "Bandit — security linter for Python [Электронный ресурс]. — URL: https://github.com/PyCQA/bandit",
        "Методические указания к курсовой работе по дисциплине «Методы и технологии программирования», 2026.",
    ]
    for i, ref in enumerate(refs, 1):
        add_para(doc, f"{i}. {ref}", indent=False)

    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЯ", 1)
    add_para(doc, "Приложение А — Руководство пользователя (README.md репозитория).", indent=False)
    add_para(doc, "Приложение Б — Скриншоты интерфейса и API (см. docs/screenshots/).", indent=False)
    add_para(doc, "Приложение В — Листинги ключевых модулей (nvd_collector.py, matcher.py).", indent=False)
    add_para(doc, "[Вставьте скриншоты по SCREENSHOTS_GUIDE.md]", bold=True)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Создан файл: {OUTPUT}")


if __name__ == "__main__":
    main()
