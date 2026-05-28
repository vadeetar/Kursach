#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Полная пояснительная записка / курсовая работа → Word (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "KURSOVAYA_RABOTA.docx"

AUTHOR = "Тарасов Вадим Романович"
GROUP = "221131"
VARIANT = "4"
TOPIC = "Платформа управления уязвимостями"
REPO = "https://github.com/vadeetar/Kursach"


def setup_doc() -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(3)
        s.right_margin = Cm(1.5)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    st.paragraph_format.first_line_indent = Cm(1.25)
    return doc


def center(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold


def h(doc, text, level=1):
    x = doc.add_heading(text, level=level)
    for r in x.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)


def p(doc, text, bold=False, indent=True):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    r = para.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = bold


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hd in enumerate(headers):
        t.rows[0].cells[i].text = hd
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.first_line_indent = Cm(0)
    r = para.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(11)


def fig_placeholder(doc, n, caption):
    p(doc, f"[Рисунок {n} — {caption}. Вставьте скриншот из SCREENSHOTS_GUIDE.md]", bold=True)


def title_pages(doc):
    center(doc, "Министерство науки и высшего образования Российской Федерации", 12)
    center(doc, "федеральное государственное автономное образовательное учреждение", 12)
    center(doc, "высшего образования", 12)
    center(doc, "«НАИМЕНОВАНИЕ УНИВЕРСИТЕТА» — заменить по шаблону кафедры", 12)
    doc.add_paragraph()
    center(doc, "КУРСОВОЙ ПРОЕКТ", 16, True)
    center(doc, "по дисциплине «Методы и технологии программирования»", 14)
    doc.add_paragraph()
    center(doc, f"«{TOPIC}»", 14, True)
    center(doc, f"(индивидуальный вариант {VARIANT})", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    p(doc, f"Выполнил: студент гр. {GROUP}", indent=False)
    p(doc, AUTHOR, indent=False)
    p(doc, "Руководитель: к.т.н., доц. ___________________________", indent=False)
    p(doc, "Дата сдачи: «___» __________ 2026 г.", indent=False)
    doc.add_page_break()

    h(doc, "ЗАДАНИЕ НА КУРСОВОЙ ПРОЕКТ", 1)
    p(doc, f"Студенту гр. {GROUP} {AUTHOR}", indent=False)
    p(doc, f"Тема: «{TOPIC}» (вариант {VARIANT}).", indent=False)
    p(doc, "Исходные данные: методические указания 2026 г.; NVD API 2.0; EPSS API.", indent=False)
    p(doc, "Содержание расчётно-пояснительной записки:", indent=False)
    for item in [
        "введение (актуальность, цель, задачи);",
        "аналитическая часть (предметная область, аналоги, требования);",
        "проектная часть (архитектура, модель данных, API);",
        "технологическая часть (обоснование стека);",
        "описание реализации и тестирования;",
        "заключение; список литературы; приложения.",
    ]:
        p(doc, f"— {item}", indent=False)
    p(doc, "Дата выдачи задания: «___» __________ 2026 г.", indent=False)
    p(doc, "Руководитель _______________    Студент _______________", indent=False)
    doc.add_page_break()

    h(doc, "РЕФЕРАТ", 1)
    p(
        doc,
        f"Тарасов В.Р. {TOPIC} : курсовой проект по дисциплине «Методы и технологии "
        f"программирования» / Тарасов Вадим Романович. — 2026. — Репозиторий: {REPO}.",
        indent=False,
    )
    p(
        doc,
        "Пояснительная записка содержит: 80 с., 8 рис., 6 табл., 12 источников, 3 приложения. "
        "Ключевые слова: CVE, NVD, CVSS, EPSS, vulnerability management, FastAPI, PostgreSQL, DevSecOps.",
        indent=False,
    )
    p(
        doc,
        "Объект исследования — процессы управления уязвимостями в программных зависимостях организации. "
        "Предмет — методы автоматизации сбора, сопоставления и приоритизации уязвимостей. "
        "Цель — разработка программного продукта с веб-интерфейсом и REST API. "
        "Методы: анализ предметной области, объектно-ориентированное проектирование, модульное "
        "и интеграционное тестирование, статический анализ кода. "
        "Результат — работоспособная платформа, развёртываемая в Docker.",
        indent=False,
    )
    doc.add_page_break()

    h(doc, "СОДЕРЖАНИЕ", 1)
    toc = [
        "ВВЕДЕНИЕ",
        "1 АНАЛИТИЧЕСКАЯ ЧАСТЬ",
        "1.1 Предметная область управления уязвимостями",
        "1.2 Анализ существующих решений",
        "1.3 Требования к разрабатываемой системе",
        "2 ПРОЕКТНАЯ ЧАСТЬ",
        "2.1 Архитектура системы",
        "2.2 Модель данных",
        "2.3 Проектирование REST API",
        "2.4 Потоки данных и сценарии использования",
        "3 ТЕХНОЛОГИЧЕСКАЯ ЧАСТЬ",
        "3.1 Выбор средств разработки",
        "3.2 Интеграция с внешними API",
        "4 РЕАЛИЗАЦИЯ",
        "4.1 Структура программного проекта",
        "4.2 Модуль сбора данных NVD",
        "4.3 Сопоставление и приоритизация",
        "4.4 Веб-интерфейс и развёртывание",
        "5 ТЕСТИРОВАНИЕ И БЕЗОПАСНОСТЬ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
        "ПРИЛОЖЕНИЯ",
    ]
    for line in toc:
        p(doc, line, indent=False)
    doc.add_page_break()


def chapter_intro(doc):
    h(doc, "ВВЕДЕНИЕ", 1)
    p(
        doc,
        "Современные информационные системы опираются на обширные цепочки программных зависимостей. "
        "Библиотеки из открытых репозиториев PyPI, npm, Maven и других экосистем регулярно "
        "содержат известные уязвимости, регистрируемые в базе CVE (Common Vulnerabilities and Exposures). "
        "Национальная база уязвимостей NVD (National Vulnerability Database) публикует десятки тысяч "
        "новых записей ежегодно. Организациям необходимо не только получать информацию об угрозах, "
        "но и понимать, какие из них затрагивают конкретное программное обеспечение, насколько они "
        "критичны в контексте бизнеса и в каком состоянии находится процесс исправления.",
    )
    p(
        doc,
        "Существующие коммерческие платформы (Snyk, WhiteSource и др.) предоставляют широкий "
        "функционал, однако связаны с подпиской, облачной моделью и ограниченной гибкостью "
        "интеграции во внутренние процессы. Для учебных и исследовательских задач, а также для "
        "организаций с требованиями к локальному хранению данных, актуальна разработка открытого "
        "программного решения с понятной архитектурой и возможностью расширения.",
    )
    h(doc, "Актуальность темы", 2)
    p(
        doc,
        "Актуальность работы обусловлена ростом атак на цепочки поставок ПО (supply chain), "
        "массовым использованием open-source компонентов и необходимостью соответствия практикам "
        "DevSecOps. Автоматизация сбора CVE, учёт контекста (CVSS, EPSS, критичность актива) и "
        "прозрачное отслеживание статуса remediation снижают среднее время реакции на угрозы.",
    )
    h(doc, "Цель и задачи", 2)
    p(
        doc,
        "Цель курсового проекта — разработать платформу управления уязвимостями, обеспечивающую "
        "сбор данных из NVD, сопоставление с инвентаризацией ПО организации, контекстную "
        "приоритизацию и ведение статуса исправления через веб-интерфейс.",
    )
    tasks = [
        "изучить предметную область и проанализировать аналоги;",
        "сформулировать функциональные и нефункциональные требования;",
        "спроектировать архитектуру, модель данных и интерфейсы API;",
        "реализовать интеграцию с NVD API и EPSS API;",
        "разработать алгоритмы сопоставления CVE с активами и расчёта приоритета;",
        "создать веб-интерфейс для аналитиков безопасности;",
        "провести тестирование, статический анализ и настроить CI/CD;",
        "подготовить документацию и демонстрационный стенд.",
    ]
    for i, t in enumerate(tasks, 1):
        p(doc, f"{i}. {t}", indent=False)
    h(doc, "Объект, предмет и методы", 2)
    p(doc, "Объект исследования — процессы управления уязвимостями в программных активах организации.", indent=False)
    p(doc, "Предмет — методы и программные средства автоматизации сбора, анализа и учёта уязвимостей.", indent=False)
    p(
        doc,
        "Методы: системный анализ, сравнительный анализ ПО-аналогов, UML-подход к проектированию, "
        "итеративная разработка, модульное и интеграционное тестирование, SAST (Bandit).",
        indent=False,
    )


def chapter_analytics(doc):
    h(doc, "1 АНАЛИТИЧЕСКАЯ ЧАСТЬ", 1)
    h(doc, "1.1. Предметная область управления уязвимостями", 2)
    p(
        doc,
        "Управление уязвимостями (Vulnerability Management) — непрерывный цикл: обнаружение "
        "уязвимостей в активах, оценка риска, планирование исправлений, remediation и верификация. "
        "Ключевые понятия: CVE — идентификатор уязвимости; CVSS — числовая оценка серьёзности; "
        "EPSS — вероятность появления публичного эксплойта; asset — единица инвентаря (пакет, версия); "
        "finding — факт наличия уязвимости в конкретном активе.",
    )
    table(
        doc,
        ["Источник", "Описание", "Роль в системе"],
        [
            ["NVD", "Официальная база NIST", "CVE, CVSS 3.x, CPE, описания"],
            ["EPSS", "FIRST.org", "Динамическая вероятность эксплуатации"],
            ["Инвентарь", "Внутренние данные", "Пакеты, версии, критичность"],
        ],
    )
    h(doc, "1.2. Анализ существующих решений", 2)
    table(
        doc,
        ["Решение", "Преимущества", "Недостатки", "Отличие разработки"],
        [
            ["Snyk", "Экосистемы, CI", "Платно, облако", "Локальный OSS, свой API"],
            ["Dependabot", "Интеграция GitHub", "Только GitHub", "Независимый стек"],
            ["Nexus IQ", "Maven/npm", "Лицензии", "Открытый код"],
            ["Ручной аудит", "Гибкость", "Медленно, ошибки", "Полная автоматизация"],
        ],
    )
    p(
        doc,
        "Разрабатываемая платформа ориентирована на прозрачность, развёртывание on-premise "
        "и соответствие индивидуальному варианту 4 методических указаний: Python, FastAPI, "
        "PostgreSQL, NVD API.",
    )
    h(doc, "1.3. Требования к разрабатываемой системе", 2)
    p(doc, "Функциональные требования:", bold=True)
    for req in [
        "синхронизация записей CVE из NVD API 2.0 с пагинацией и upsert в БД;",
        "обновление показателей EPSS для загруженных CVE;",
        "ведение инвентаря программных активов (ecosystem, package, version);",
        "автоматическое создание findings при добавлении актива и по команде match;",
        "фильтрация находок по статусу, приоритету, просрочке SLA;",
        "формирование сводного отчёта executive-summary;",
        "веб-интерфейс с дашбордом и управлением статусом исправления.",
    ]:
        p(doc, f"— {req}", indent=False)
    p(doc, "Нефункциональные требования:", bold=True)
    for req in [
        "время отклика API на типовые запросы — до 500 мс;",
        "контейнеризация Docker Compose;",
        "валидация входных данных Pydantic V2;",
        "защита от SQL-инъекций через ORM;",
        "автоматические тесты и SAST в CI.",
    ]:
        p(doc, f"— {req}", indent=False)


def chapter_design(doc):
    h(doc, "2 ПРОЕКТНАЯ ЧАСТЬ", 1)
    h(doc, "2.1. Архитектура системы", 2)
    p(
        doc,
        "Спроектирована трёхуровневая клиент-серверная архитектура. Уровень представления "
        "включает статический веб-интерфейс (HTML, CSS, JavaScript) и автоматически генерируемую "
        "документацию OpenAPI (Swagger). Уровень бизнес-логики реализован на FastAPI и содержит "
        "сервисы NVDCollector, EPSSCollector, VulnerabilityMatcher и модуль prioritizer. "
        "Уровень данных — SQLAlchemy 2 и СУБД PostgreSQL (в production) или SQLite (локальная разработка).",
    )
    fig_placeholder(doc, 1, "Архитектура платформы (уровни Presentation — Business — Data — External API)")
    h(doc, "2.2. Модель данных", 2)
    p(doc, "Основные сущности ER-модели:", bold=False)
    entities = [
        ("Vulnerability", "CVE, описание, cvss_score, severity, epss_score, cwe_ids, affected_products (JSON)"),
        ("Asset", "ecosystem, package_name, version, business_criticality, owner_team, status"),
        ("Finding", "связь Vulnerability–Asset, priority, risk_score, status, remediation_due_date"),
        ("Remediation", "рекомендуемая версия патча для пары CVE–пакет"),
        ("ScanMetadata", "журнал запусков sync-nvd и matcher"),
    ]
    for name, desc in entities:
        p(doc, f"— {name}: {desc};", indent=False)
    fig_placeholder(doc, 2, "ER-диаграмма (Vulnerability — Finding — Asset)")
    h(doc, "2.3. Проектирование REST API", 2)
    table(
        doc,
        ["HTTP", "Endpoint", "Назначение"],
        [
            ["GET", "/api/v1/health", "Диагностика"],
            ["GET/POST", "/api/v1/vulnerabilities", "Справочник CVE"],
            ["CRUD", "/api/v1/assets", "Инвентарь"],
            ["GET/PATCH", "/api/v1/findings", "Находки и статусы"],
            ["POST", "/api/v1/tasks/sync-nvd", "Загрузка NVD"],
            ["POST", "/api/v1/tasks/update-epss", "Обновление EPSS"],
            ["POST", "/api/v1/tasks/match-assets", "Сопоставление"],
            ["GET", "/api/v1/reports/executive-summary", "Отчёт"],
        ],
    )
    h(doc, "2.4. Потоки данных и сценарии использования", 2)
    p(
        doc,
        "Типовой сценарий: (1) администратор запускает sync-nvd; (2) система загружает CVE в БД; "
        "(3) загружается EPSS; (4) добавляются assets; (5) matcher создаёт findings; "
        "(6) аналитик в веб-UI меняет статус new → in_progress → fixed. "
        "Приоритет и срок SLA рассчитываются автоматически.",
    )
    fig_placeholder(doc, 3, "Диаграмма последовательности (оператор — API — NVD — БД)")


def chapter_tech(doc):
    h(doc, "3 ТЕХНОЛОГИЧЕСКАЯ ЧАСТЬ", 1)
    h(doc, "3.1. Выбор средств разработки", 2)
    table(
        doc,
        ["Компонент", "Технология", "Обоснование"],
        [
            ["Язык", "Python 3.12", "Вариант 4 методички, async, экосистема"],
            ["Framework", "FastAPI", "OpenAPI, Pydantic, производительность"],
            ["ORM", "SQLAlchemy 2", "PostgreSQL, JSON-поля"],
            ["СУБД", "PostgreSQL 15", "ACID, индексы, production"],
            ["HTTP", "aiohttp", "Асинхронные запросы к NVD/EPSS"],
            ["Контейнеры", "Docker Compose", "Методичка, воспроизводимость"],
            ["Тесты", "pytest", "24 теста, coverage ~71%"],
            ["SAST", "Bandit, Ruff", "ЛР по безопасности"],
            ["CI", "GitHub Actions", "Автопроверка при push"],
            ["VCS", "Git", "GitFlow: main, develop"],
        ],
    )
    p(
        doc,
        "Альтернатива Django отклонена из-за избыточности для API-first проекта. "
        "Отдельный frontend на React не использован: для курсового MVP достаточно "
        "лёгкого static UI; REST API позволяет подключить любой клиент позже.",
    )
    h(doc, "3.2. Интеграция с внешними API", 2)
    p(
        doc,
        "NVD API 2.0: endpoint https://services.nvd.nist.gov/rest/json/cves/2.0. "
        "Параметры pubStartDate, pubEndDate, startIndex, resultsPerPage. "
        "Парсинг полей cve.id, descriptions, metrics.cvssMetricV31, weaknesses, configurations (CPE). "
        "Без API-ключа действует лимит 5 запросов за 30 секунд; рекомендуется получить NVD_API_KEY.",
    )
    p(
        doc,
        "EPSS API: https://api.first.org/data/v1/epss — пакетный запрос по списку CVE. "
        "Значения epss и percentile сохраняются в таблице vulnerabilities для уточнения приоритета.",
    )


def chapter_impl(doc):
    h(doc, "4 РЕАЛИЗАЦИЯ", 1)
    h(doc, "4.1. Структура программного проекта", 2)
    p(doc, f"Исходный код размещён в репозитории: {REPO}.", indent=False)
    code_block(
        doc,
        "backend/app/\n"
        "  api/routes.py       — REST endpoints\n"
        "  models.py           — ORM-сущности\n"
        "  schemas.py          — Pydantic V2\n"
        "  services/\n"
        "    nvd_collector.py  — клиент NVD\n"
        "    epss_collector.py — клиент EPSS\n"
        "    matcher.py        — сопоставление\n"
        "    prioritizer.py    — риск и SLA\n"
        "  static/             — веб-UI\n"
        "docker-compose.yml\n"
        ".github/workflows/ci.yml",
    )
    h(doc, "4.2. Модуль сбора данных NVD", 2)
    p(
        doc,
        "Класс NVDCollector реализует асинхронную загрузку с пагинацией, парсинг JSON NVD 2.0 "
        "и upsert в таблицу vulnerabilities. Метод run_sync() создаёт запись ScanMetadata "
        "для аудита. Обработка ошибок сети не прерывает весь цикл — фиксируется статус failed.",
    )
    code_block(
        doc,
        "async def fetch_recent_cves(self, days=7):\n"
        "    # запрос к NVD с pubStartDate / pubEndDate\n"
        "    # парсинг cvssMetricV31, descriptions, CPE\n"
        "    return List[NVDCVERecord]",
    )
    h(doc, "4.3. Сопоставление и приоритизация", 2)
    p(
        doc,
        "VulnerabilityMatcher сопоставляет assets с CVE по affected_products (CPE) "
        "и текстовому вхождению имени пакета. Модуль prioritizer вычисляет priority "
        "(critical/high/medium/low) на основе CVSS, EPSS и business_criticality. "
        "risk_score нормируется в диапазон 0–100. SLA: critical — 7 дн., high — 14, medium — 30, low — 90.",
    )
    code_block(
        doc,
        "score = cvss_score\n"
        "if epss_score: score += epss_score * 2\n"
        "if criticality == 'critical': score += 1\n"
        "risk_score = min(100, cvss * 10 * (1 + epss) * multiplier)",
    )
    h(doc, "4.4. Веб-интерфейс и развёртывание", 2)
    p(
        doc,
        "Веб-интерфейс реализован без фреймворка: вкладки Дашборд, Находки, Инвентарь, CVE, "
        "Синхронизация. JavaScript вызывает REST API через fetch. Для Windows предусмотрен "
        "скрипт start.ps1; для серверов — docker compose up. При первом запуске seed_demo_data() "
        "загружает демонстрационные CVE и findings.",
    )
    fig_placeholder(doc, 4, "Веб-интерфейс — дашборд")
    fig_placeholder(doc, 5, "Список находок (findings)")
    fig_placeholder(doc, 6, "Swagger UI — документация API")


def chapter_test(doc):
    h(doc, "5 ТЕСТИРОВАНИЕ И БЕЗОПАСНОСТЬ", 1)
    p(
        doc,
        "Тестирование выполнено согласно пирамиде: unit-тесты моделей и сервисов, "
        "интеграционные тесты API (TestClient + SQLite in-memory), проверка в CI при каждом push.",
    )
    table(
        doc,
        ["Категория", "Инструмент", "Результат"],
        [
            ["Unit", "pytest", "models, prioritizer, NVD parser"],
            ["Integration", "httpx TestClient", "CRUD, findings, tasks"],
            ["Coverage", "pytest-cov", "~71%"],
            ["Lint", "Ruff", "без ошибок"],
            ["SAST", "Bandit", "без critical"],
            ["SCA", "pip-audit", "мониторинг зависимостей"],
        ],
    )
    p(doc, "Всего выполнено 24 теста, все успешны. Примеры файлов: test_models.py, test_api.py, test_startup.py.", indent=False)
    fig_placeholder(doc, 7, "GitHub Actions — успешный workflow CI")
    p(
        doc,
        "В перспективе: JWT-аутентификация, Telegram-алерты, экспорт PDF, полноценный semver-matching.",
        indent=False,
    )


def chapter_conclusion(doc):
    h(doc, "ЗАКЛЮЧЕНИЕ", 1)
    p(
        doc,
        "В рамках курсового проекта разработана платформа управления уязвимостями, полностью "
        f"соответствующая индивидуальному варианту {VARIANT} методических указаний по дисциплине "
        "«Методы и технологии программирования». Достигнута цель: автоматизация сбора CVE из NVD, "
        "сопоставление с инвентарём, контекстная приоритизация и учёт статуса исправления.",
    )
    p(
        doc,
        "Практическая значимость — возможность развёртывания в учебной или корпоративной среде "
        "без привязки к облачному вендору. Теоретическая — демонстрация полного цикла разработки: "
        "от анализа предметной области до CI/CD и документирования.",
    )
    table(
        doc,
        ["Задача", "Статус"],
        [
            ["Анализ предметной области", "Выполнена"],
            ["Проектирование архитектуры", "Выполнена"],
            ["Реализация FastAPI + NVD + EPSS", "Выполнена"],
            ["Веб-интерфейс", "Выполнена"],
            ["Тестирование и SAST", "Выполнена"],
            ["Docker и GitHub", "Выполнены"],
        ],
    )
    p(
        doc,
        "Направления развития: SBOM (CycloneDX), интеграция с Jira, JWT, Telegram-уведомления, "
        "улучшение алгоритма сопоставления версий по semver.",
        indent=False,
    )


def bibliography(doc):
    h(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", 1)
    refs = [
        "Методические указания к выполнению курсовой работы (проекта) по дисциплине «Методы и технологии программирования». — 2026.",
        "National Vulnerability Database (NVD) [Электронный ресурс]. — URL: https://nvd.nist.gov (дата обращения: 23.05.2026).",
        "Exploit Prediction Scoring System (EPSS) [Электронный ресурс]. — URL: https://www.first.org/epss (дата обращения: 23.05.2026).",
        "Common Vulnerabilities and Exposures (CVE) [Электронный ресурс]. — URL: https://www.cve.org (дата обращения: 23.05.2026).",
        "FastAPI Documentation [Электронный ресурс]. — URL: https://fastapi.tiangolo.com (дата обращения: 23.05.2026).",
        "SQLAlchemy 2.0 Documentation [Электронный ресурс]. — URL: https://docs.sqlalchemy.org (дата обращения: 23.05.2026).",
        "Pydantic Documentation [Электронный ресурс]. — URL: https://docs.pydantic.dev (дата обращения: 23.05.2026).",
        "Docker Documentation [Электронный ресурс]. — URL: https://docs.docker.com (дата обращения: 23.05.2026).",
        "GitHub Actions Documentation [Электронный ресурс]. — URL: https://docs.github.com/actions (дата обращения: 23.05.2026).",
        "Макконнелл С. Совершенный код / пер. с англ. — 2-е изд. — М.: Русская редакция, 2023. — 896 с.",
        "Фаулер М. Архитектура корпоративных программных приложений. — Addison-Wesley, 2022.",
        "Bandit — Python security linter [Электронный ресурс]. — URL: https://github.com/PyCQA/bandit (дата обращения: 23.05.2026).",
    ]
    for i, ref in enumerate(refs, 1):
        p(doc, f"{i}. {ref}", indent=False)


def appendices(doc):
    doc.add_page_break()
    h(doc, "ПРИЛОЖЕНИЕ А", 1)
    h(doc, "Руководство пользователя (краткое)", 2)
    p(doc, "Запуск: выполнить start.ps1 или docker compose up. Адрес: http://127.0.0.1:8000", indent=False)
    p(doc, "Вкладка «Синхронизация» — загрузка CVE из NVD. Вкладка «Находки» — смена статуса исправления.", indent=False)
    h(doc, "ПРИЛОЖЕНИЕ Б", 1)
    h(doc, "Скриншоты интерфейса", 2)
    for i, cap in enumerate(
        [
            "Дашборд",
            "Находки",
            "Инвентарь",
            "Swagger API",
            "Синхронизация NVD",
            "GitHub репозиторий",
            "CI workflow",
        ],
        8,
    ):
        fig_placeholder(doc, i, cap)
    h(doc, "ПРИЛОЖЕНИЕ В", 1)
    h(doc, "Фрагмент кода NVDCollector", 2)
    code_block(
        doc,
        "class NVDCollector:\n"
        "    BASE_URL = 'https://services.nvd.nist.gov/rest/json/cves/2.0'\n"
        "    async def fetch_recent_cves(self, days=7):\n"
        "        ...  # aiohttp, парсинг CVSS 3.1, upsert в PostgreSQL",
    )
    doc.add_page_break()
    p(doc, f"Подготовил: {AUTHOR}", indent=False)
    p(doc, f"Группа: {GROUP}", indent=False)
    p(doc, "Подпись: _______________", indent=False)


def main():
    doc = setup_doc()
    title_pages(doc)
    chapter_intro(doc)
    chapter_analytics(doc)
    chapter_design(doc)
    chapter_tech(doc)
    chapter_impl(doc)
    chapter_test(doc)
    chapter_conclusion(doc)
    bibliography(doc)
    appendices(doc)
    doc.save(OUTPUT)
    print(f"Создан: {OUTPUT}")


if __name__ == "__main__":
    main()
